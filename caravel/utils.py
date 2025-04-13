###############################################################################
# NSAp - Copyright (C) CEA, 2019
# Distributed under the terms of the CeCILL-B license, as published by
# the CEA-CNRS-INRIA. Refer to the LICENSE file or to
# http://www.cecill.info/licences/Licence_CeCILL-B_V1-en.html
# for details.
###############################################################################


"""
This module provides common utilities.
"""

# Imports
import functools
import json
import os
import re
import time
from datetime import date, datetime

from dateutil.relativedelta import relativedelta
from docx import Document
from tqdm import tqdm


def export_report(report, timestamp, outfile):
    """Export the report in docx format.

    Parameters
    ----------
    report: str
        a dictionary with recursive keys and values.
    timestamp: str
        a timestamp.
    outfile: str
        the path to the generated docx file.
    """
    document = Document()
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = f"NeuroSpin\tReporting\t{timestamp}"
    paragraph = document.add_paragraph(
        f"\n\n\n\nYou will find below the report generated on '{timestamp}'. "
        "If you have any questions please use the contact mail: rlink@cea.fr."
    )
    for family, family_item in report.items():
        document.add_heading(family.replace(".", " ").title())
        for validator, validator_item in family_item.items():
            split_validator = re.findall(r"[A-Z][^A-Z]*", validator)
            document.add_heading(" ".join(split_validator))
            paragraph = document.add_paragraph(
                "\n\n Below the table summarizing the errors.\n\n"
            )
            for key, values in validator_item.items():
                table = document.add_table(rows=len(values), cols=2)
                cell = table.cell(0, 0)
                cell.text = key
                for idx, val in enumerate(values):
                    cell = table.cell(idx, 1)
                    cell.text = val
    document.save(outfile)


def get_logs_to_remove(log_dir, cut_date=None):
    """Return all files that are older than cut_date
    according to their name.

    Parameters
    ----------
    log_dir: str
        path to the directory to clean.
    cut_date: str or None
        date from before which files are returned. If None,
        return files than are older than one year old.
        Should be formatted %Y-%m-%d (e.g. 2024-08-06).
    """

    print(f"Clean {log_dir}")
    if cut_date is None:
        current_date = date.today()
        cut_date = current_date - relativedelta(years=1)
    else:
        cut_date = datetime.strptime(cut_date, "%Y-%m-%d").date()
    print("Cut date", cut_date)

    files2remove = []

    for filename in os.listdir(log_dir):
        date_regex = r"20\d\d-(0|1)\d-(0|1|2|3)\d"
        file_date = re.search(date_regex, filename)
        if file_date is not None:
            file_date = datetime.strptime(file_date.group(0), "%Y-%m-%d").date()
            if file_date <= cut_date:
                files2remove.append(filename)

    return files2remove


def clean_logs_dir(log_dir, cut_date=None):
    """Remove all files that are older than cut_date.

    Parameters
    ----------
    log_dir: str
        path to the directory to clean.
    cut_date:
        date from before which files are suppressed. If None,
        remove files than are older than one year old.
        Should be formatted %Y-%m-%d (e.g. 2024-08-06).
    """

    filenames = get_logs_to_remove(log_dir, cut_date=cut_date)
    print(f"Number of files to remove: {len(filenames)}")
    for filename in tqdm(filenames):
        filepath = os.path.join(log_dir, filename)
        os.remove(filepath)


def monitor(func):
    """A decorator to monitor function and log its status in a root directory.
    The input function parameters must be set via the 'CARAVEL_ROOT'
    and 'CARAVEL_NAME' environment variables.
    """
    root = os.environ.get("CARAVEL_ROOT", None)
    name = os.environ.get("CARAVEL_NAME", None)
    is_monitor = root is not None and name is not None
    if is_monitor:
        assert os.path.isdir(root), root

    @functools.wraps(func)
    def decorated(*args, **kwargs):
        try:
            tic = time.time()
            res = func(*args, **kwargs)
            toc = time.time()
            info = {"status": "healthy", "duration": toc - tic}
            if is_monitor:
                with open(os.path.join(root, f"{name}.json"), "wt") as of:
                    json.dump(info, of, indent=4)
            return res
        except Exception as e:
            info = {"status": "dead"}
            if is_monitor:
                with open(os.path.join(root, f"{name}.json"), "wt") as of:
                    json.dump(info, of, indent=4)
            raise e

    return decorated
